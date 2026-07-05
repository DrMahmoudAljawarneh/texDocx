import os, subprocess, tempfile
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from PIL import Image as PILImage

JATS_NS = "http://www.ncbi.nlm.nih.gov/JATS1"
MATH_NS = "http://www.w3.org/1998/Math/MathML"


# ── Improvement #1: Document-level defaults ──────────────────────────────────

def _setup_document_defaults(doc):
    """Configure the document with professional academic defaults."""
    # Page margins: 1 inch all around
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Default body font: Times New Roman, 12pt, single spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Title style (level 0)
    title_style = doc.styles['Title']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(6)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Times New Roman'
    h4.font.size = Pt(12)
    h4.font.bold = False
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(2)


def generate_docx(task_id: str, xml_content: str, output_path: Path, citation_style: str = "ieee", algorithm_render: str = "text", asset_dir: Path = None, template_path: Path = None):
    from tasks.celery_app import publish_log

    job_dir = output_path.parent

    if template_path and template_path.exists():
        publish_log(task_id, "Info", f"Using journal template: {template_path.name}")
        doc = Document(str(template_path))
        # Clear existing paragraphs if any, preserving styles
        for paragraph in doc.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
    else:
        doc = Document()
        _setup_document_defaults(doc)  # Improvement #1

    def _safe_add_heading(text, level):
        try:
            return doc.add_heading(text, level=level)
        except KeyError:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(16 if level == 0 else max(10, 14 - level))
            return p

    root = etree.fromstring(xml_content.encode("utf-8"))
    ns = {"jats": JATS_NS, "xlink": "http://www.w3.org/1999/xlink", "ltx": "http://dlmf.nist.gov/LaTeXML"}

    CSL_DIR = Path(__file__).parent.parent / "csl"
    VALID_STYLES = {"ieee", "apa", "mla", "chicago", "harvard"}

    csl_style_name = citation_style.lower()
    if csl_style_name not in VALID_STYLES:
        csl_style_name = "ieee"

    is_numeric_style = csl_style_name == "ieee"
    is_author_year = csl_style_name in ("apa", "harvard", "chicago", "mla")

    def format_author_year_inline(keys, csl_map):
        parts = []
        for k in keys:
            entry = csl_map.get(k)
            if not entry:
                parts.append(k)
                continue
            authors = entry.get("author", [])
            year = ""
            issued = entry.get("issued", {})
            if issued:
                dp = issued.get("date-parts", [])
                if dp and dp[0]:
                    year = str(dp[0][0])
            if not authors:
                parts.append(f"({year})" if year else k)
                continue
            family = authors[0].get("family", "")
            if csl_style_name == "mla":
                parts.append(family if family else k)
            elif csl_style_name == "chicago":
                parts.append(f"{family} {year}" if family and year else k)
            else:
                parts.append(f"{family}, {year}" if family and year else k)
        if csl_style_name == "mla":
            return f"({', '.join(parts)})"
        return f"({'; '.join(parts)})"

    LATEXML_HOST = os.environ.get("LATEXML_HOST", "latexml")
    LATEXML_PORTS = os.environ.get("LATEXML_PORTS", "3334").split(",")

    def render_algo_png(algo_snippet):
        import urllib.request, urllib.parse
        import re
        # Force [H] placement on algorithm environments to prevent floating
        snippet_fixed = re.sub(r'\\begin\{(algorithm|algorithm\*)\}', r'\\begin{\1}[H]', algo_snippet)
        # Detect which algorithmic package the snippet uses
        if re.search(r'\\(STATE|WHILE|ENDWHILE|IF|ELSIF|ELSE|ENDIF|FOR|ENDFOR|REQUIRE|ENSURE|COMMENT|AND|OR|XOR|LOOP|ENDLOOP|REPEAT|UNTIL)', snippet_fixed):
            algo_pkg = r"""\usepackage{algorithmic}"""
        else:
            algo_pkg = r"""\usepackage{algpseudocode}"""
        doc = r"""\documentclass{article}
\usepackage[paperwidth=1000pt,paperheight=5000pt,margin=10pt]{geometry}
\usepackage{amsmath}
\usepackage{algorithm}
%s
\usepackage{float}
\usepackage{varwidth}
\pagestyle{empty}
\begin{document}
\begin{varwidth}{980pt}
%s
\end{varwidth}
\end{document}""" % (algo_pkg, snippet_fixed)
        data = urllib.parse.urlencode({"source": doc}).encode()
        last_exc = ""
        for port in LATEXML_PORTS:
            try:
                port = port.strip()
                req = urllib.request.Request(
                    f"http://{LATEXML_HOST}:{port}/render-png",
                    data=data,
                    headers={"Content-Type": "application/x-www-form-urlencoded"}
                )
                with urllib.request.urlopen(req, timeout=60) as resp:
                    if resp.status == 200:
                        png_data = resp.read()
                        try:
                            import subprocess
                            cmd = ["convert", "-limit", "memory", "2GB", "-limit", "map", "2GB", "-limit", "disk", "4GB", "png:-", "-trim", "+repage", "png8:-"]
                            res = subprocess.run(cmd, input=png_data, capture_output=True, check=True)
                            return res.stdout
                        except subprocess.CalledProcessError as e:
                            err_msg = e.stderr.decode('utf-8', errors='replace')
                            publish_log(task_id, "Warning", f"convert -trim failed: {e}\n{err_msg}")
                            return png_data
                        except Exception as e:
                            publish_log(task_id, "Warning", f"convert -trim failed (other): {e}")
                            return png_data
                    else:
                        last_exc = f"HTTP {resp.status}"
            except Exception as e:
                last_exc = str(e)
                continue
        publish_log(task_id, "Warning", f"render_algo_png failed: {last_exc}")
        return None

    refs_json = job_dir / "references.json"
    csl_data = None
    csl_map = {}
    if refs_json.exists():
        try:
            import json
            csl_data = json.loads(refs_json.read_text(encoding="utf-8"))
            csl_map = {item.get("id"): item for item in csl_data}
        except Exception:
            pass

    def find_any(parent_el, tags):
        for t in tags:
            el = parent_el.find(t, ns)
            if el is not None:
                return el
        return None

    def find_all_any(parent_el, tags):
        for t in tags:
            els = parent_el.findall(t, ns)
            if els:
                return els
        return []

    def extract_title_or_caption(el, is_caption=False):
        tag_el = find_any(el, ["ltx:tag", "jats:tag", "tag"])
        if tag_el is not None:
            tag_text = "".join(tag_el.itertext()).strip()
            text_parts = []
            if el.text:
                text_parts.append(el.text)
            for ch in el:
                ch_tag = etree.QName(ch).localname
                if ch_tag not in ("tag", "tags"):
                    text_parts.append("".join(ch.itertext()))
                if ch.tail:
                    text_parts.append(ch.tail)
            main_text = "".join(text_parts).strip()
            if tag_text:
                if is_caption:
                    return f"{tag_text}: {main_text}"
                else:
                    return f"{tag_text} {main_text}"
            else:
                return main_text
        else:
            return "".join(el.itertext()).strip()

    # ── Build citation ref_map ────────────────────────────────────────────
    ref_map = {}
    ref_index = 1
    for ref_elem in find_all_any(root, [".//ltx:ref", ".//jats:ref", ".//ref"]):
        ref_id = ref_elem.get("id")
        if ref_id:
            clean_id = ref_id.replace("ref-", "")
            ref_map[clean_id] = str(ref_index)
            ref_map[ref_id] = str(ref_index)
            ref_index += 1

    csl_bibliography = None
    if csl_data and refs_json.exists():
        try:
            import warnings
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                from citeproc import CitationStylesStyle, CitationStylesBibliography
                from citeproc import Citation, CitationItem
                from citeproc.source.json import CiteProcJSON

                csl_file = CSL_DIR / f"{csl_style_name}.csl"
                if csl_file.exists():
                    style = CitationStylesStyle(str(csl_file))
                    bib_source = CiteProcJSON(csl_data)
                    csl_bibliography = CitationStylesBibliography(style, bib_source)
                    for item in csl_data:
                        key = item.get("id")
                        if key:
                            csl_bibliography.register(Citation([CitationItem(key)]))
        except Exception as e:
            publish_log(task_id, "Warning", f"citeproc init failed: {e}")
            csl_bibliography = None

    # ── Render document title ─────────────────────────────────────────────
    title_el = find_any(root, [".//ltx:title", ".//jats:article-title", ".//article-title"])
    if title_el is not None:
        parent = title_el.getparent()
        parent_tag = etree.QName(parent).localname if parent is not None else ""
        if parent_tag not in ("sec", "section", "subsection", "subsubsection", "ref-list", "figure", "fig", "table-wrap"):
            title_text = "".join(title_el.itertext()).strip()
            if title_text:
                _safe_add_heading(title_text, level=0)

    # ── Improvement #8: Author / affiliation metadata ─────────────────────
    _render_authors(root, doc, ns, find_any, find_all_any)

    # ── Image resolution helper (Improvement #5) ─────────────────────────
    def resolve_image(href):
        if not asset_dir:
            return None
        # Strip vector extensions to resolve converted png/jpg first
        base_href = href
        for ext in (".pdf", ".eps", ".svg", ".PDF", ".EPS", ".SVG"):
            if href.endswith(ext):
                base_href = href[:-len(ext)]
                break

        for ext in (".png", ".jpg", ".jpeg", ".gif", ".PNG", ".JPG", ".JPEG", ".GIF"):
            candidate = asset_dir / (base_href + ext)
            if candidate.exists():
                return candidate

        candidate = asset_dir / href
        if candidate.exists():
            return candidate
        return None

    def _smart_image_width(img_path, max_width_inches=5.5):
        """Return an appropriate width that never upscales the image beyond its native size."""
        try:
            with PILImage.open(str(img_path)) as img:
                w_px, h_px = img.size
                dpi = img.info.get("dpi", (150, 150))
                dpi_x = dpi[0] if isinstance(dpi, (tuple, list)) else dpi
                if dpi_x <= 0:
                    dpi_x = 150
                native_width_inches = w_px / dpi_x
                return Inches(min(native_width_inches, max_width_inches))
        except Exception:
            return Inches(max_width_inches)

    algo_counter = [0]

    # ── Main recursive walker ─────────────────────────────────────────────
    def process_children(parent_elem, doc_parent):
        for child in parent_elem:
            if not isinstance(child.tag, str):
                continue
            tag = etree.QName(child).localname

            # ── Sections / headings ───────────────────────────────────────
            if tag in ("sec", "section", "subsection", "subsubsection"):
                heading_el = find_any(child, ["ltx:title", "jats:title", "title"])
                level = _infer_level(child)
                if heading_el is not None:
                    heading_text = extract_title_or_caption(heading_el, is_caption=False)
                    if heading_text:
                        _safe_add_heading(heading_text, level=level)
                process_children(child, doc)

            # ── Paragraphs ────────────────────────────────────────────────
            elif tag == "p":
                if algorithm_render == "image":
                    # Check if this paragraph contains algorithm content
                    has_caption = False
                    has_algo_markers = False
                    for c in child:
                        if not isinstance(c.tag, str):
                            continue
                        ct = etree.QName(c).localname
                        if ct == "text" and "caption" in (c.get("class") or ""):
                            has_caption = True
                        if ct == "ERROR":
                            t = c.text or ""
                            if any(cmd in t for cmd in ("\\Require", "\\State", "\\For", "\\If", "\\While", "\\Repeat", "\\Ensure", "\\Return", "\\Procedure", "\\Function")):
                                has_algo_markers = True
                    if has_algo_markers:
                        snippet_path = job_dir / f"algo_{algo_counter[0]}.tex"
                        algo_counter[0] += 1
                        if snippet_path.exists():
                            snippet = snippet_path.read_text(encoding="utf-8")
                            publish_log(task_id, "Info", f"Rendering algorithm {algo_counter[0]-1} via latexml PNG")
                            png_data = render_algo_png(snippet)
                            if png_data:
                                publish_log(task_id, "Info", f"Algorithm PNG rendered ({len(png_data)} bytes)")
                                tmp_png = job_dir / f"_algo_{algo_counter[0]-1}.png"
                                tmp_png.write_bytes(png_data)
                                para = doc.add_paragraph()
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = para.add_run()
                                run.add_picture(str(tmp_png), width=Inches(5.5))
                                tmp_png.unlink()
                                continue
                            else:
                                publish_log(task_id, "Warning", f"Algorithm PNG rendering failed, falling back to text")

                para = doc.add_paragraph()
                _add_inline_content(child, para)

            # ── Improvement #7: Abstract ──────────────────────────────────
            elif tag == "abstract":
                _safe_add_heading("Abstract", level=1)
                # Process abstract children; render paragraphs in italic
                for abs_child in child:
                    if not isinstance(abs_child.tag, str):
                        continue
                    abs_tag = etree.QName(abs_child).localname
                    if abs_tag == "p":
                        para = doc.add_paragraph()
                        _add_inline_content(abs_child, para)
                        for run in para.runs:
                            run.italic = True
                    elif abs_tag in ("sec", "section"):
                        process_children(abs_child, doc)
                    else:
                        process_children(abs_child, doc)
                # Add visual separator after abstract
                sep = doc.add_paragraph()
                sep.paragraph_format.space_before = Pt(6)
                sep.paragraph_format.space_after = Pt(6)

            # ── Table wraps ───────────────────────────────────────────────
            elif tag == "table-wrap":
                caption_el = find_any(child, [".//ltx:caption", ".//jats:caption", ".//caption"])
                caption_text = ""
                if caption_el is not None:
                    caption_text = extract_title_or_caption(caption_el, is_caption=True)
                if caption_text:
                    cap_para = doc.add_paragraph()
                    run = cap_para.add_run(caption_text)
                    run.bold = True
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # #6: center captions

                table = child.find(".//jats:table", ns) or child.find(".//table")
                if table is not None:
                    _render_table(table, doc)

            elif tag in ("title", "label", "caption", "toccaption"):
                continue

            elif tag == "alternatives":
                display = child.get("display", "inline")
                if display == "block":
                    block_para = doc.add_paragraph()
                    block_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    block_para.paragraph_format.space_before = Pt(12)
                    block_para.paragraph_format.space_after = Pt(6)
                    math_el = child.find(f".//{{{MATH_NS}}}math")
                    if math_el is not None:
                        _insert_equation_image(math_el, block_para, is_block=True)
                    else:
                        process_children(child, doc)
                else:
                    process_children(child, doc)

            # ── Figures ───────────────────────────────────────────────────
            elif tag in ("fig", "figure"):
                caption_el = find_any(child, [".//ltx:caption", ".//jats:caption", ".//caption"])
                caption_text = ""
                if caption_el is not None:
                    caption_text = extract_title_or_caption(caption_el, is_caption=True)
                else:
                    toc_caption = find_any(child, [".//ltx:toccaption", ".//jats:toccaption", ".//toccaption"])
                    if toc_caption is not None:
                        caption_text = extract_title_or_caption(toc_caption, is_caption=True)

                graphic = find_any(child, [".//ltx:graphics", ".//ltx:graphic", ".//jats:graphics", ".//jats:graphic", ".//graphics", ".//graphic"])
                if graphic is not None:
                    href = graphic.get("{%s}href" % ns["xlink"]) or graphic.get("href") or graphic.get("graphic") or ""
                    img_path = resolve_image(href)
                    if img_path:
                        try:
                            # Improvement #6: Center figures with breathing room
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            img_para.paragraph_format.space_before = Pt(12)
                            run = img_para.add_run()
                            # Improvement #5: Smart image sizing
                            width = _smart_image_width(img_path)
                            run.add_picture(str(img_path), width=width)

                            if caption_text:
                                cap_para = doc.add_paragraph()
                                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # #6: center caption
                                cap_run = cap_para.add_run(caption_text)
                                cap_run.italic = True
                                cap_run.font.size = Pt(10)
                                cap_para.paragraph_format.space_after = Pt(12)
                        except Exception as e:
                            publish_log(task_id, "Warning", f"Failed to insert image {href}: {e}")

            # ── References section ────────────────────────────────────────
            elif tag == "ref-list":
                process_children(child, doc)
            elif tag == "ref":
                para = doc.add_paragraph(style='List Number' if 'List Number' in doc.styles else None)
                if csl_bibliography is not None:
                    try:
                        citems = []
                        for entry in csl_data:
                            citems.append(entry)
                        formatted = list(csl_bibliography.bibliography())
                        if formatted:
                            for entry_text in formatted:
                                para.add_run(str(entry_text))
                                break
                            else:
                                _add_inline_content(child, para)
                        else:
                            _add_inline_content(child, para)
                        continue
                    except Exception:
                        pass
                citation = find_any(child, [".//ltx:element-citation", ".//jats:element-citation", ".//element-citation"])
                mixed = find_any(child, [".//ltx:mixed-citation", ".//jats:mixed-citation", ".//mixed-citation"])
                if citation is not None:
                    parts = []
                    authors = find_all_any(citation, [".//ltx:string-name", ".//jats:string-name", ".//string-name"])
                    author_names = []
                    for author in authors:
                        surname = find_any(author, [".//ltx:surname", ".//jats:surname", ".//surname"])
                        given = find_any(author, [".//ltx:given-names", ".//jats:given-names", ".//given-names"])
                        s_text = surname.text if surname is not None else ""
                        g_text = given.text if given is not None else ""
                        if s_text and g_text:
                            author_names.append(f"{s_text}, {g_text}")
                        elif s_text:
                            author_names.append(s_text)
                    if author_names:
                        parts.append(", ".join(author_names))

                    title = find_any(citation, [".//ltx:title", ".//jats:title", ".//title", ".//ltx:article-title", ".//jats:article-title", ".//article-title"])
                    if title is not None and title.text:
                        parts.append(f'"{title.text}"')

                    source = find_any(citation, [".//ltx:source", ".//jats:source", ".//source", ".//ltx:journal", ".//jats:journal", ".//journal"])
                    if source is not None and source.text:
                        parts.append(source.text)

                    year = find_any(citation, [".//ltx:year", ".//jats:year", ".//year"])
                    if year is not None and year.text:
                        parts.append(year.text)

                    if not parts:
                        text = "".join(citation.itertext()).strip()
                        if text:
                            parts.append(text)

                    para.add_run(". ".join(parts))
                elif mixed is not None:
                    _add_inline_content(mixed, para)
                else:
                    _add_inline_content(child, para)

            elif tag == "mixed-citation":
                para = doc.add_paragraph()
                _add_inline_content(child, para)

            # ── Improvement #12: Page break before References ─────────────
            elif tag in ("back", "bibliography"):
                # Insert a page break before References
                page_break_para = doc.add_paragraph()
                run = page_break_para.add_run()
                run.add_break(WD_BREAK.PAGE)
                _safe_add_heading("References", level=1)
                process_children(child, doc)

            # ── Algorithm floats ──────────────────────────────────────────
            elif tag == "float":
                cls = child.get("class") or ""
                if "algorithm" in cls:
                    if algorithm_render == "image":
                        snippet_path = job_dir / f"algo_{algo_counter[0]}.tex"
                        algo_counter[0] += 1
                        if snippet_path.exists():
                            snippet = snippet_path.read_text(encoding="utf-8")
                            publish_log(task_id, "Info", f"Rendering algorithm {algo_counter[0]-1} via latexml PNG ({len(snippet)} bytes)")
                            png_data = render_algo_png(snippet)
                            if png_data:
                                publish_log(task_id, "Info", f"Algorithm PNG rendered ({len(png_data)} bytes)")
                                tmp_png = job_dir / f"_algo_{algo_counter[0]-1}.png"
                                tmp_png.write_bytes(png_data)
                                para = doc.add_paragraph()
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = para.add_run()
                                run.add_picture(str(tmp_png), width=Inches(5.5))
                                tmp_png.unlink()
                                continue
                            else:
                                publish_log(task_id, "Warning", f"Algorithm PNG rendering failed, falling back to text")
                        else:
                            publish_log(task_id, "Warning", f"Algorithm snippet not found: {snippet_path}")

                    caption_el = find_any(child, [".//ltx:caption", ".//jats:caption", ".//caption"])
                    caption_text = ""
                    if caption_el is not None:
                        caption_text = extract_title_or_caption(caption_el, is_caption=True)

                    listinglines = child.findall(".//ltx:listingline", ns) + child.findall(".//jats:listingline", ns) + child.findall(".//listingline")
                    if not listinglines:
                        process_children(child, doc)
                        continue

                    table = doc.add_table(rows=2, cols=1)

                    tbl = table._tbl
                    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

                    tblBorders = OxmlElement('w:tblBorders')
                    for side, sz in [('top', '8'), ('bottom', '8'), ('left', '4'), ('right', '4')]:
                        el = OxmlElement(f'w:{side}')
                        el.set(qn('w:val'), 'single')
                        el.set(qn('w:sz'), sz)
                        el.set(qn('w:space'), '0')
                        el.set(qn('w:color'), '333333')
                        tblBorders.append(el)
                    for side in ['insideH', 'insideV']:
                        el = OxmlElement(f'w:{side}')
                        el.set(qn('w:val'), 'none')
                        el.set(qn('w:sz'), '0')
                        el.set(qn('w:space'), '0')
                        el.set(qn('w:color'), 'auto')
                        tblBorders.append(el)
                    tblPr.append(tblBorders)

                    cap_cell = table.cell(0, 0)
                    cap_para = cap_cell.paragraphs[0]
                    if caption_text:
                        cap_run = cap_para.add_run(caption_text)
                        cap_run.bold = True
                        cap_run.font.name = 'Times New Roman'
                        cap_run.font.size = Pt(11)

                    cap_tcPr = cap_cell._tc.get_or_add_tcPr()
                    cap_borders = OxmlElement('w:tcBorders')
                    cap_bottom = OxmlElement('w:bottom')
                    cap_bottom.set(qn('w:val'), 'single')
                    cap_bottom.set(qn('w:sz'), '4')
                    cap_bottom.set(qn('w:space'), '0')
                    cap_bottom.set(qn('w:color'), '999999')
                    cap_borders.append(cap_bottom)
                    cap_tcPr.append(cap_borders)

                    code_cell = table.cell(1, 0)
                    first_line = True
                    indent_level = 0
                    import re
                    for line in listinglines:
                        line_text = "".join(line.itertext()).lower()
                        words = re.findall(r'\b\w+\b', line_text)
                        
                        if any(w in words for w in ["end", "until", "else", "elsif", "endwhile", "endif", "endfor"]):
                            indent_level = max(0, indent_level - 1)

                        if first_line:
                            para = code_cell.paragraphs[0]
                            first_line = False
                        else:
                            para = code_cell.add_paragraph()

                        para.paragraph_format.left_indent = Inches(0.15 + indent_level * 0.2)
                        para.paragraph_format.space_before = Pt(1)
                        para.paragraph_format.space_after = Pt(1)
                        para.paragraph_format.line_spacing = 1.0

                        if any(w in words for w in ["while", "for", "if", "procedure", "function", "repeat", "else", "elsif", "loop"]):
                            indent_level += 1

                        tag_el = find_any(line, [".//ltx:tag", ".//jats:tag", ".//tag"])
                        line_num = ""
                        if tag_el is not None and tag_el.text:
                            line_num = tag_el.text.strip()
                        if line_num:
                            run_num = para.add_run(f"{line_num}:  ")
                            run_num.font.name = "Courier New"
                            run_num.font.size = Pt(9)
                            run_num.font.color.rgb = RGBColor(120, 120, 120)

                        _add_inline_content(line, para)
                        for r in para.runs:
                            r.font.name = "Courier New"
                            r.font.size = Pt(9)
                else:
                    process_children(child, doc)

            elif tag == "listing":
                process_children(child, doc)
            elif tag == "listingline":
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                tag_el = find_any(child, [".//ltx:tag", ".//jats:tag", ".//tag"])
                line_num = ""
                if tag_el is not None and tag_el.text:
                    line_num = tag_el.text.strip()
                if line_num:
                    run_num = para.add_run(f"{line_num}  ")
                    run_num.font.name = "Courier New"
                _add_inline_content(child, para)
                for r in para.runs:
                    r.font.name = "Courier New"

            # ── Improvement #10: List handling ────────────────────────────
            elif tag in ("list", "itemize", "enumerate"):
                list_type = child.get("list-type") or child.get("class") or ""
                is_ordered = ("ordered" in list_type or "enumerate" in list_type or tag == "enumerate")
                for item in child:
                    if not isinstance(item.tag, str):
                        continue
                    item_tag = etree.QName(item).localname
                    if item_tag in ("item", "list-item"):
                        if is_ordered:
                            style_name = 'List Number' if 'List Number' in doc.styles else None
                        else:
                            style_name = 'List Bullet' if 'List Bullet' in doc.styles else None
                        para = doc.add_paragraph(style=style_name)
                        # Collect <p> elements: they may be direct children
                        # or nested inside <para> wrappers
                        p_elements = _collect_p_elements(item)
                        if p_elements:
                            _add_inline_content(p_elements[0], para)
                            for extra_p in p_elements[1:]:
                                extra_para = doc.add_paragraph()
                                extra_para.paragraph_format.left_indent = Inches(0.5)
                                _add_inline_content(extra_p, extra_para)
                        else:
                            _add_inline_content(item, para)
                    else:
                        # Could be a nested list or other element
                        process_children(item, doc)

            # ── Improvement #9: Footnotes ─────────────────────────────────
            elif tag == "note":
                role = child.get("role") or child.get("class") or ""
                if "foot" in role:
                    # Render as an endnote-style paragraph (true footnotes require
                    # deep OxmlElement work; for now use a styled paragraph)
                    note_text = "".join(child.itertext()).strip()
                    if note_text:
                        para = doc.add_paragraph()
                        run = para.add_run(note_text)
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(80, 80, 80)
                        para.paragraph_format.left_indent = Inches(0.5)
                else:
                    process_children(child, doc)

            else:
                process_children(child, doc)

    def _infer_level(sec_elem):
        parent = sec_elem.getparent()
        level = 1
        while parent is not None:
            parent_tag = etree.QName(parent).localname
            if parent_tag in ("sec", "section", "subsection", "subsubsection"):
                level += 1
            parent = parent.getparent()
        return min(level, 4)

    def _collect_p_elements(item_elem):
        """Collect <p> elements from a list <item>, handling the LaTeXML
        structure where <p> is nested inside <para> wrappers.
        Structure: <item> → <tags> + <para> → <p>
        """
        p_elements = []
        for ch in item_elem:
            if not isinstance(ch.tag, str):
                continue
            ch_tag = etree.QName(ch).localname
            if ch_tag == "p":
                p_elements.append(ch)
            elif ch_tag == "para":
                # LaTeXML wraps <p> inside <para>
                for sub in ch:
                    if isinstance(sub.tag, str) and etree.QName(sub).localname == "p":
                        p_elements.append(sub)
            # Skip <tags> and other non-content elements
        return p_elements

    # ── Improvement #2: Fixed inline content — preserve whitespace ────────

    def _add_inline_content(elem, para):
        """Recursively add inline content to a paragraph, preserving inter-word spaces."""
        # Use raw text (do NOT strip) to preserve spaces between elements
        text = elem.text or ""
        if text:
            para.add_run(text)
        for child in elem:
            if not isinstance(child.tag, str):
                continue
            tag = etree.QName(child).localname
            if tag in ("tags", "tag"):
                # Skip tag elements but DO preserve their tail
                tail = child.tail or ""
                if tail:
                    para.add_run(tail)
                continue

            if tag in ("bold", "b", "strong"):
                child_text = child.text or ""
                if child_text:
                    run = para.add_run(child_text)
                    run.bold = True
                # Recurse into children of the bold element
                for sub in child:
                    if isinstance(sub.tag, str):
                        _add_inline_child(sub, para, bold=True)
                    sub_tail = sub.tail or ""
                    if sub_tail:
                        r = para.add_run(sub_tail)
                        r.bold = True
            elif tag in ("italic", "i", "em"):
                child_text = child.text or ""
                if child_text:
                    run = para.add_run(child_text)
                    run.italic = True
                for sub in child:
                    if isinstance(sub.tag, str):
                        _add_inline_child(sub, para, italic=True)
                    sub_tail = sub.tail or ""
                    if sub_tail:
                        r = para.add_run(sub_tail)
                        r.italic = True
            elif tag in ("monospace", "code", "tt"):
                child_text = child.text or ""
                if child_text:
                    run = para.add_run(child_text)
                    run.font.name = "Courier New"
                for sub in child:
                    if isinstance(sub.tag, str):
                        _add_inline_child(sub, para, mono=True)
                    sub_tail = sub.tail or ""
                    if sub_tail:
                        r = para.add_run(sub_tail)
                        r.font.name = "Courier New"

            # ── Improvement #11: Cross-references — no underline ──────
            elif tag == "xref":
                xref_text = "".join(child.itertext())
                if xref_text:
                    para.add_run(xref_text)

            # ── Improvement #4: Citation resolution ──────────────────
            elif tag == "bibref":
                bibrefs_attr = child.get("bibrefs") or ""
                keys = [k.strip() for k in bibrefs_attr.split(",") if k.strip()]
                if is_author_year and csl_map:
                    resolved_text = format_author_year_inline(keys, csl_map)
                    para.add_run(resolved_text)
                else:
                    labels = []
                    for k in keys:
                        val = ref_map.get(k) or ref_map.get("ref-" + k)
                        if val:
                            labels.append(val)
                        else:
                            labels.append(k)
                    resolved_text = ", ".join(labels)
                    if resolved_text:
                        para.add_run(resolved_text)
                    else:
                        txt = "".join(child.itertext())
                        if txt:
                            para.add_run(txt)

            elif tag == "alternatives":
                math_el = child.find(f".//{{{MATH_NS}}}math")
                if math_el is not None:
                    _insert_equation_image(math_el, para)
                else:
                    _add_inline_content(child, para)

            elif tag == "math":
                _insert_equation_image(child, para)
            elif tag == "Math":
                _insert_equation_image(child, para)

            elif tag == "svg":
                continue

            elif tag in ("graphic", "inline-graphic", "graphics"):
                href = child.get("{%s}href" % ns["xlink"]) or child.get("href") or child.get("graphic") or ""
                img_path = resolve_image(href)
                if img_path:
                    try:
                        run = para.add_run()
                        width = _smart_image_width(img_path, max_width_inches=4.0)
                        run.add_picture(str(img_path), width=width)
                    except Exception as e:
                        publish_log(task_id, "Warning", f"Failed to insert inline image {href}: {e}")

            elif tag == "break":
                # Line break
                run = para.add_run()
                run.add_break()

            elif tag == "sup":
                sup_text = "".join(child.itertext())
                if sup_text:
                    run = para.add_run(sup_text)
                    run.font.superscript = True

            elif tag == "sub":
                sub_text = "".join(child.itertext())
                if sub_text:
                    run = para.add_run(sub_text)
                    run.font.subscript = True

            # ── LaTeXML <text font="bold/italic"> elements ────────────
            elif tag == "text":
                font_attr = child.get("font") or ""
                child_text = child.text or ""
                if child_text:
                    run = para.add_run(child_text)
                    if "bold" in font_attr:
                        run.bold = True
                    if "italic" in font_attr:
                        run.italic = True
                    if "monospace" in font_attr or "typewriter" in font_attr:
                        run.font.name = "Courier New"
                # Recurse into sub-elements within the text element
                for sub in child:
                    if isinstance(sub.tag, str):
                        _add_inline_content(sub, para)
                    sub_tail = sub.tail or ""
                    if sub_tail:
                        r = para.add_run(sub_tail)
                        if "bold" in font_attr:
                            r.bold = True
                        if "italic" in font_attr:
                            r.italic = True

            # ── LaTeXML <para> wrapper — just recurse ─────────────────
            elif tag == "para":
                _add_inline_content(child, para)

            # ── <cite> — handle bracket suppression for author-year styles ──
            elif tag == "cite":
                if is_author_year:
                    cite_text = child.text or ""
                    if cite_text.strip() == "[":
                        pass
                    elif cite_text:
                        para.add_run(cite_text)
                    for sub in child:
                        if not isinstance(sub.tag, str):
                            continue
                        sub_tag = etree.QName(sub).localname
                        if sub_tag == "bibref":
                            bibrefs_attr = sub.get("bibrefs") or ""
                            keys = [k.strip() for k in bibrefs_attr.split(",") if k.strip()]
                            resolved = format_author_year_inline(keys, csl_map)
                            para.add_run(resolved)
                            sub_tail = sub.tail or ""
                            if sub_tail.strip() != "]":
                                if sub_tail:
                                    para.add_run(sub_tail)
                        else:
                            _add_inline_child(sub, para)
                            sub_tail = sub.tail or ""
                            if sub_tail:
                                para.add_run(sub_tail)
                else:
                    _add_inline_content(child, para)

            else:
                # Unknown inline element — recurse into its children
                _add_inline_content(child, para)

            # Preserve tail text (the text after the closing tag of this child)
            tail = child.tail or ""
            if tail:
                para.add_run(tail)

    def _add_inline_child(elem, para, bold=False, italic=False, mono=False):
        """Helper for nested inline formatting (e.g. bold inside italic)."""
        tag = etree.QName(elem).localname
        if tag in ("math", "Math"):
            _insert_equation_image(elem, para)
            return
        text = "".join(elem.itertext())
        if text:
            run = para.add_run(text)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if mono:
                run.font.name = "Courier New"

    _eq_counter = 0

    def _render_equation_svg(math_elem):
        """Render a math element to SVG bytes. Handles:
        - <math> inside <alternatives> with SVG sibling
        - Bare <Math> (LaTeXML) elements via latex2mathml + ziamath
        Returns (svg_bytes, is_block) or (None, False) on failure.
        """
        parent = math_elem.getparent()
        if parent is not None and etree.QName(parent).localname == "alternatives":
            for sibling in parent:
                sibling_tag = etree.QName(sibling).localname
                if sibling_tag == "svg":
                    return (etree.tostring(sibling, encoding="utf-8"), parent.get("display", "inline") == "block")
        tex_str = math_elem.get("tex")
        if tex_str:
            try:
                from latex2mathml.converter import convert as latex2mathml_convert
                import ziamath as zm
                mathml_str = latex2mathml_convert(tex_str)
                eqn = zm.Math(mathml_str)
                svg_str = eqn.svg()
                display = (math_elem.get("mode") or math_elem.get("display") or "inline") == "display"
                return (svg_str.encode("utf-8"), display)
            except Exception as e:
                publish_log(task_id, "Warning", f"On-the-fly equation rendering failed: {e}")
        return (None, False)

    def _insert_equation_image(math_elem, para, is_block=False):
        """Render equation to PNG and insert as picture in the paragraph."""
        nonlocal _eq_counter
        svg_bytes, detected_block = _render_equation_svg(math_elem)
        if svg_bytes is None:
            run = para.add_run("[eq]")
            run.italic = True
            return
        _eq_counter += 1
        svg_path = job_dir / f"_eqn{_eq_counter}.svg"
        png_path = svg_path.with_suffix(".png")
        try:
            svg_path.write_bytes(svg_bytes)
            subprocess.run(
                ["magick", str(svg_path), "-background", "white", "-flatten", "-density", "200", str(png_path)],
                capture_output=True, timeout=30, check=True,
            )
            with PILImage.open(str(png_path)) as img:
                w, h = img.size
                max_w = 5.5
                aspect = h / w if w > 0 else 1
                display_width = Inches(min(w / 150, max_w))
                display_height = Inches(display_width.inches * aspect)
            run = para.add_run()
            run.add_picture(str(png_path), width=display_width)
        except Exception as e:
            publish_log(task_id, "Warning", f"Equation image insertion failed: {e}")
            run = para.add_run("[eq]")
            run.italic = True
        finally:
            try:
                svg_path.unlink(missing_ok=True)
            except Exception:
                pass
            try:
                png_path.unlink(missing_ok=True)
            except Exception:
                pass

    def _insert_math(math_elem, para):
        _insert_equation_image(math_elem, para, is_block=False)

    # ── Improvement #3: Table rendering with colspan/rowspan + rich cells ─

    def _render_table(table_elem, doc):
        rows_el = table_elem.findall(".//tr") or table_elem.findall(".//jats:tr", ns)
        if not rows_el:
            return

        # First pass: determine grid dimensions
        num_cols = 0
        row_data = []
        for row in rows_el:
            cells = row.findall("td") + row.findall("th") + row.findall("jats:td", ns) + row.findall("jats:th", ns)
            col_count = 0
            for cell in cells:
                colspan = int(cell.get("colspan", 1))
                col_count += colspan
            num_cols = max(num_cols, col_count)
            row_data.append(cells)

        if num_cols == 0:
            return

        doc_table = doc.add_table(rows=len(rows_el), cols=num_cols)
        doc_table.style = "Table Grid"
        doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Track merged cells for rowspan
        occupied = {}  # (row, col) -> True for cells occupied by rowspan

        for ri, cells in enumerate(row_data):
            ci = 0
            for cell in cells:
                # Skip occupied cells (from previous rowspan)
                while (ri, ci) in occupied:
                    ci += 1
                if ci >= num_cols:
                    break

                colspan = int(cell.get("colspan", 1))
                rowspan = int(cell.get("rowspan", 1))
                is_header = etree.QName(cell).localname == "th"

                # Merge cells for colspan
                if colspan > 1:
                    end_ci = min(ci + colspan - 1, num_cols - 1)
                    try:
                        doc_table.cell(ri, ci).merge(doc_table.cell(ri, end_ci))
                    except Exception:
                        pass

                # Merge cells for rowspan
                if rowspan > 1:
                    end_ri = min(ri + rowspan - 1, len(rows_el) - 1)
                    try:
                        doc_table.cell(ri, ci).merge(doc_table.cell(end_ri, ci))
                    except Exception:
                        pass
                    # Mark the spanned cells as occupied
                    for rr in range(ri + 1, end_ri + 1):
                        for cc in range(ci, min(ci + colspan, num_cols)):
                            occupied[(rr, cc)] = True

                # Render cell content with rich inline support
                doc_cell = doc_table.cell(ri, ci)
                # Clear default paragraph
                if doc_cell.paragraphs:
                    p = doc_cell.paragraphs[0]
                    p.clear()
                else:
                    p = doc_cell.add_paragraph()

                _add_inline_content(cell, p)

                # Bold header cells
                if is_header:
                    for run in p.runs:
                        run.bold = True

                ci += colspan

    # ── Improvement #8: Author rendering ──────────────────────────────────

    process_children(root, doc)
    doc.save(str(output_path))
    publish_log(task_id, "Info", f"DOCX saved to {output_path}")


def _render_authors(root, doc, ns, find_any, find_all_any):
    """Render author names and affiliations below the title."""
    # Try LaTeXML creator elements
    creators = find_all_any(root, [
        ".//ltx:creator", ".//jats:contrib", ".//contrib",
    ])
    if not creators:
        return

    author_names = []
    affiliations = []

    for creator in creators:
        # LaTeXML: <creator><personname>...</personname><contact>...</contact></creator>
        name_el = find_any(creator, [
            "ltx:personname", "jats:name", "name", "personname",
        ])
        if name_el is not None:
            name_text = "".join(name_el.itertext()).strip()
            if name_text:
                author_names.append(name_text)

        contact_el = find_any(creator, [
            "ltx:contact", "jats:aff", "aff", "contact",
        ])
        if contact_el is not None:
            aff_text = "".join(contact_el.itertext()).strip()
            if aff_text and aff_text not in affiliations:
                affiliations.append(aff_text)

    if author_names:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(", ".join(author_names))
        run.font.size = Pt(12)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)

    if affiliations:
        for aff in affiliations:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(aff)
            run.font.size = Pt(10)
            run.italic = True
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
